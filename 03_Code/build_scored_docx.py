"""
build_scored_docx.py

Creates one professional, human-readable Word document per provider,
showing every prompt with its response(s), per-criterion scores WITH
justifications, computed Success/Failure outcomes, and total scores - the
scored equivalent of Week 4's build_results_docx.py.

This is the document that answers "on what basis was this response scored
this way": every criterion is listed by name with its individual 0-5 score,
a one-sentence justification, and whether it counted as Success or Failure
against the documented threshold. For the 20 stochastic-sample prompts, ALL
5 scored runs are shown together so score variance across runs is directly
visible.

Reads from 05_Logs_Results/Scored_Results/<Provider>_Scores/ directly (does
not require build_scored_tables.py to have been run first).

OUTPUT FILES (written to 05_Logs_Results/Scored_Results/Readable_Scores/):
  - gemini_scores.docx
  - mistral_scores.docx
  - groq_scores.docx

Requires: pip install python-docx

Run with the VS Code Run button:
    python 03_Code/build_scored_docx.py
"""

import json
import os
import re
import glob

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.join(os.path.dirname(__file__), "..")
SCORED_RESULTS_DIR = os.path.join(BASE_DIR, "05_Logs_Results", "Scored_Results")
OUTPUT_DIR = os.path.join(SCORED_RESULTS_DIR, "Readable_Scores")

PROVIDERS = {
    "gemini": {
        "score_dir": os.path.join(SCORED_RESULTS_DIR, "Gemini_Scores"),
        "file_prefix": "gemini_",
        "output_name": "gemini_scores.docx",
        "title": "Gemini - Scored Results",
    },
    "mistral": {
        "score_dir": os.path.join(SCORED_RESULTS_DIR, "Mistral_Scores"),
        "file_prefix": "mistral_",
        "output_name": "mistral_scores.docx",
        "title": "Mistral - Scored Results",
    },
    "groq": {
        "score_dir": os.path.join(SCORED_RESULTS_DIR, "Groq_Scores"),
        "file_prefix": "groq_",
        "output_name": "groq_scores.docx",
        "title": "Groq - Scored Results",
    },
}

FILENAME_PATTERN = re.compile(r"^(?P<prefix>\w+_)(?P<id>P\d{3})(?:_run(?P<run>\d+))?_scored\.json$")

COLUMNS = [
    ("ID", 0.5),
    ("Category", 1.0),
    ("Difficulty", 0.6),
    ("Prompt", 1.6),
    ("Criterion Scores, Justification & Outcome", 2.8),
    ("Total", 0.5),
    ("Outcome", 0.7),
    ("Latency (s)", 0.7),
    ("Run", 0.4),
]


def _prompt_id_sort_key(prompt_id):
    match = re.search(r"\d+", prompt_id)
    return int(match.group()) if match else 0


def _load_and_group(score_dir, prefix):
    pattern = os.path.join(score_dir, f"{prefix}*_scored.json")
    files = [f for f in glob.glob(pattern) if not f.endswith(".tmp")]

    from collections import defaultdict
    grouped = defaultdict(list)
    for filepath in files:
        filename = os.path.basename(filepath)
        match = FILENAME_PATTERN.match(filename)
        if not match:
            continue
        with open(filepath, "r", encoding="utf-8") as f:
            try:
                grouped[match.group("id")].append(json.load(f))
            except json.JSONDecodeError:
                print(f"  WARNING: could not parse {filepath}, skipping.")

    for prompt_id in grouped:
        grouped[prompt_id].sort(key=lambda r: r.get("run_number", 1))

    return grouped


def _format_criterion_scores(criterion_scores) -> str:
    """Every criterion's name, score, outcome, and justification together -
    this is the actual 'on what basis' documentation the supervisor asked for."""
    lines = []
    for c in criterion_scores:
        name = c.get("name", "?")
        score = c.get("score", "?")
        outcome = c.get("outcome", "?")
        justification = c.get("justification", "")
        lines.append(f"{name}: {score}/5 [{outcome}]")
        if justification:
            lines.append(f"   {justification}")
    return "\n".join(lines) if lines else "(no criterion scores recorded)"


def _set_cell_shading(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _set_repeat_header(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def _add_multiline_text(paragraph, text, font_size, bold=False, color=None):
    lines = str(text).split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            paragraph.add_run().add_break(WD_BREAK.LINE)
        run = paragraph.add_run(line)
        run.font.size = Pt(font_size)
        run.bold = bold
        if color:
            run.font.color.rgb = color


def build_doc_for_provider(name, config):
    grouped = _load_and_group(config["score_dir"], config["file_prefix"])
    prompt_ids = sorted(grouped.keys(), key=_prompt_id_sort_key)

    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width, section.page_height = new_width, new_height
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(config["title"])
    run.bold = True
    run.font.size = Pt(18)

    total_rows = sum(len(runs) for runs in grouped.values())
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(
        f"{len(prompt_ids)} prompts, {total_rows} scored runs - "
        f"Success threshold: score >= 3/5 per criterion"
    )
    sub_run.italic = True
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    table = doc.add_table(rows=1, cols=len(COLUMNS))
    table.style = "Table Grid"
    table.autofit = False

    header_row = table.rows[0]
    _set_repeat_header(header_row)
    for i, (col_name, width) in enumerate(COLUMNS):
        cell = header_row.cells[i]
        cell.width = Inches(width)
        p = cell.paragraphs[0]
        r = p.add_run(col_name)
        r.bold = True
        r.font.size = Pt(9)
        _set_cell_shading(cell, "D9E2F3")

    for prompt_id in prompt_ids:
        runs = grouped[prompt_id]
        is_multi_run = len(runs) > 1

        for run_record in runs:
            row = table.add_row()

            total_score = run_record.get("total_score", "")
            max_score = run_record.get("max_score", "")
            overall_outcome = run_record.get("overall_outcome", "")
            criterion_text = _format_criterion_scores(run_record.get("criterion_scores", []))
            run_number = run_record.get("run_number", 1)

            values = [
                prompt_id if run_number == 1 else "",  # avoid repeating ID on every run row
                run_record.get("category", ""),
                run_record.get("difficulty", ""),
                run_record.get("prompt", "") if run_number == 1 else "",
                criterion_text,
                f"{total_score}/{max_score}",
                overall_outcome,
                str(run_record.get("latency_seconds", "-")),
                str(run_number) if is_multi_run else "-",
            ]
            for i, (col_name, width) in enumerate(COLUMNS):
                cell = row.cells[i]
                cell.width = Inches(width)
                p = cell.paragraphs[0]
                if col_name == "Outcome":
                    color = RGBColor(0x1E, 0x7E, 0x34) if overall_outcome == "Success" else RGBColor(0xB0, 0x00, 0x00)
                    _add_multiline_text(p, values[i], font_size=8.5, bold=True, color=color)
                else:
                    _add_multiline_text(p, values[i], font_size=8)

            if is_multi_run:
                for cell_idx in [0, 1, 2, 3]:
                    _set_cell_shading(row.cells[cell_idx], "FFF3CD")
            if overall_outcome == "Failure":
                _set_cell_shading(row.cells[6], "FCE4E4")

    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        "Scoring methodology: each response scored by an LLM judge (openrouter/free) against "
        "the SPECIFIC evaluation criteria for that exact prompt, using standardized 0-5 anchors "
        "(0=Complete Failure, 3=Minimum Success Threshold, 5=Complete Success). Each criterion's "
        "Success/Failure outcome and the total score are computed programmatically from the "
        "judge's numeric scores, not trusted from the judge's own labeling. A prompt's overall "
        "outcome is Success only if every one of its criteria individually succeeded. "
        "Stochastic-sample prompts (highlighted) show all 5 independently scored runs."
    )
    footer_run.italic = True
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    output_path = os.path.join(OUTPUT_DIR, config["output_name"])
    doc.save(output_path)

    print(f"[{name}] {len(prompt_ids)} prompts ({total_rows} scored rows) written to {output_path}")
    return len(prompt_ids)


def main():
    print(f"Building human-readable scored tables in: {os.path.abspath(OUTPUT_DIR)}\n")
    for name, config in PROVIDERS.items():
        build_doc_for_provider(name, config)
    print("\nDone.")


if __name__ == "__main__":
    main()

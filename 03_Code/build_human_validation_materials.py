"""
build_human_validation_materials.py

Generates the two materials needed for Week 6's Human-in-the-Loop
validation: a read-only reference document showing the sampled
prompts/responses, and a blank scoring template for you to fill in your
own scores.

DELIBERATELY BLIND: the reading document and scoring template show the
prompt, the model's actual response, and the evaluation criteria - but NOT
the LLM judge's scores. This is intentional: scoring while seeing the
judge's existing scores would bias your judgment toward agreeing with it,
defeating the purpose of an independent validation check.

Reads the 20-prompt sample from 04_Datasets/human_validation_sample.json
(built by select_human_validation_sample.py) and pulls each provider's
actual response from 05_Logs_Results/Combined_Results/<provider>_prompts_
and_results.json (Week 4's raw consolidated data).

OUTPUT:
  - 05_Logs_Results/Human_Validation/Human_Validation_Reading.docx
      Read-only reference: prompt + response + criteria, for all 60 items
      (20 prompts x 3 providers), organized by provider.
  - 05_Logs_Results/Human_Validation/human_scores_template.json
      Blank template - fill in your own score (0-5) and justification for
      each criterion of each of the 60 items.

Run with the VS Code Run button:
    python 03_Code/build_human_validation_materials.py
"""

import json
import os
import re

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.join(os.path.dirname(__file__), "..")
SAMPLE_PATH = os.path.join(BASE_DIR, "04_Datasets", "human_validation_sample.json")
COMBINED_RESULTS_DIR = os.path.join(BASE_DIR, "05_Logs_Results", "Combined_Results")
OUTPUT_DIR = os.path.join(BASE_DIR, "05_Logs_Results", "Survey")

PROVIDERS = ["gemini", "mistral", "groq"]

# Number of independent human raters. Each gets their OWN separate scoring
# template and their responses are saved in their own file - never merged -
# so individual rater scores remain traceable, and inter-rater agreement
# (how much raters agree with EACH OTHER, not just with the judge) can be
# computed, per the proposal's Statistical Analysis section which lists
# "Inter-rater agreement" as a required metric.
NUM_RATERS = 3


def parse_criteria(evaluation_criteria: str):
    criteria = []
    for line in str(evaluation_criteria).strip().split("\n"):
        line = line.strip()
        if not line:
            continue
        match = re.match(r"^(.+?):\s*0-5\s*\((.+)\)$", line)
        if match:
            criteria.append({"name": match.group(1).strip(), "description": match.group(2).strip()})
        else:
            criteria.append({"name": line, "description": ""})
    return criteria


def load_sample():
    with open(SAMPLE_PATH, "r", encoding="utf-8") as f:
        return json.load(f)


def load_provider_responses(provider):
    """Loads Week 4's raw consolidated data for one provider, indexed by prompt ID."""
    path = os.path.join(COMBINED_RESULTS_DIR, f"{provider}_prompts_and_results.json")
    if not os.path.exists(path):
        print(f"  WARNING: {path} not found - has build_results_tables.py been run?")
        return {}
    with open(path, "r", encoding="utf-8") as f:
        records = json.load(f)
    return {r["id"]: r for r in records}


def _set_cell_shading(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _add_multiline(paragraph, text, size=10, bold=False):
    lines = str(text).split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            paragraph.add_run().add_break(WD_BREAK.LINE)
        run = paragraph.add_run(line)
        run.font.size = Pt(size)
        run.bold = bold


def build_reading_document(sample, provider_data):
    doc = Document()
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_w, new_h = section.page_height, section.page_width
    section.page_width, section.page_height = new_w, new_h
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Week 6 Human Validation \u2014 Reading Reference")
    r.bold = True
    r.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = subtitle.add_run(
        f"{sample['total_prompts_sampled']} prompts x 3 providers = "
        f"{sample['total_responses_for_validation']} responses. "
        f"Judge scores intentionally NOT shown - score independently using "
        f"human_scores_template.json."
    )
    sr.italic = True
    sr.font.size = Pt(10)
    sr.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    for provider in PROVIDERS:
        heading = doc.add_paragraph()
        hr = heading.add_run(provider.capitalize())
        hr.bold = True
        hr.font.size = Pt(14)
        hr.font.color.rgb = RGBColor(0x2E, 0x5C, 0x8A)

        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        table.autofit = False
        widths = [0.6, 1.8, 3.5, 3.5]
        header_cells = table.rows[0].cells
        for i, (label, w) in enumerate(zip(["ID", "Category / Difficulty", "Prompt", "Response"], widths)):
            header_cells[i].width = Inches(w)
            p = header_cells[i].paragraphs[0]
            hr2 = p.add_run(label)
            hr2.bold = True
            hr2.font.size = Pt(9)
            _set_cell_shading(header_cells[i], "D9E2F3")

        responses = provider_data.get(provider, {})
        for item in sample["prompts"]:
            pid = item["id"]
            record = responses.get(pid)
            response_text = "(response not found - check Combined_Results)"
            if record and record.get("runs"):
                response_text = record["runs"][0].get("response_text", response_text)

            row = table.add_row()
            values = [pid, f"{item['category']}\n{item['difficulty']}", item["prompt"], response_text]
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
                p = row.cells[i].paragraphs[0]
                _add_multiline(p, values[i], size=8.5)

        doc.add_paragraph()

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    output_path = os.path.join(OUTPUT_DIR, "Human_Validation_Reading.docx")
    doc.save(output_path)
    print(f"Reading document written: {output_path}")


def build_scoring_template(sample, provider_data, rater_id):
    template = {
        "rater_id": rater_id,
        "instructions": (
            f"RATER {rater_id}: For each item below, read the prompt and response in "
            "Human_Validation_Reading.docx, then fill in YOUR OWN score "
            "(integer 0-5) for each criterion, using the same scale as "
            "rubric.docx: 0=Complete Failure, 1=Major Failure, "
            "2=Partial Failure, 3=Minimum Success (threshold), "
            "4=Strong Success, 5=Complete Success. Do NOT look at the "
            "LLM judge's existing scores, and do NOT look at other raters' "
            "scores, while doing this - score independently. This is YOUR "
            "OWN file; do not share it with other raters until everyone "
            "has finished, to keep scoring genuinely independent."
        ),
        "items": [],
    }

    for provider in PROVIDERS:
        responses = provider_data.get(provider, {})
        for item in sample["prompts"]:
            pid = item["id"]
            criteria = parse_criteria(item["evaluation_criteria"])
            template["items"].append({
                "id": pid,
                "provider": provider,
                "category": item["category"],
                "difficulty": item["difficulty"],
                "your_scores": [
                    {"criterion_name": c["name"], "your_score": None, "your_justification": ""}
                    for c in criteria
                ],
            })

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    output_path = os.path.join(OUTPUT_DIR, f"human_scores_rater{rater_id}.json")
    with open(output_path, "w", encoding="utf-8") as f:
        json.dump(template, f, indent=2, ensure_ascii=False)
    print(f"Scoring template written for Rater {rater_id}: {output_path} ({len(template['items'])} items)")


def main():
    sample = load_sample()
    provider_data = {p: load_provider_responses(p) for p in PROVIDERS}

    build_reading_document(sample, provider_data)
    for rater_id in range(1, NUM_RATERS + 1):
        build_scoring_template(sample, provider_data, rater_id)

    print(f"\n{NUM_RATERS} separate rater template(s) created - each rater fills in "
          f"their OWN file (human_scores_rater1.json, human_scores_rater2.json, ...), "
          f"independently, without seeing each other's or the judge's scores.")
    print("Next step: once all raters have filled in their files, run compare_human_vs_judge.py.")


if __name__ == "__main__":
    main()

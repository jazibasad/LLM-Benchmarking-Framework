"""
build_statistical_report_docx.py

Converts the two JSON outputs from Week 7 (reliability_statistics.json and
statistical_analysis.json) into ONE polished, human-readable Word document
with proper tables - for direct inclusion in your report or for reviewing
results without reading raw JSON.

Reads:
  - 05_Logs_Results/Statistical_Analysis/reliability_statistics.json
  - 05_Logs_Results/Statistical_Analysis/statistical_analysis.json

OUTPUT: 05_Logs_Results/Statistical_Analysis/Week7_Statistical_Report.docx

Requires: pip install python-docx

Run with the VS Code Run button:
    python 03_Code/build_statistical_report_docx.py
"""

import json
import os

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.join(os.path.dirname(__file__), "..")
STATS_DIR = os.path.join(BASE_DIR, "05_Logs_Results", "Statistical_Analysis")
RELIABILITY_PATH = os.path.join(STATS_DIR, "reliability_statistics.json")
ANALYSIS_PATH = os.path.join(STATS_DIR, "statistical_analysis.json")
OUTPUT_PATH = os.path.join(STATS_DIR, "Week7_Statistical_Report.docx")

GREEN = RGBColor(0x1E, 0x7E, 0x34)
RED = RGBColor(0xB0, 0x00, 0x00)
GRAY = RGBColor(0x66, 0x66, 0x66)


def _load(path):
    if not os.path.exists(path):
        return None
    with open(path, "r", encoding="utf-8") as f:
        return json.load(f)


def _set_cell_shading(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.autofit = True

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        r = p.add_run(h)
        r.bold = True
        r.font.size = Pt(9)
        _set_cell_shading(hdr_cells[i], "D9E2F3")
        if col_widths:
            hdr_cells[i].width = Inches(col_widths[i])

    for row_data in rows:
        row = table.add_row()
        for i, (value, highlight) in enumerate(row_data):
            cell = row.cells[i]
            p = cell.paragraphs[0]
            r = p.add_run(str(value))
            r.font.size = Pt(9)
            if highlight == "green":
                r.bold = True
                r.font.color.rgb = GREEN
            elif highlight == "red":
                r.bold = True
                r.font.color.rgb = RED
            if col_widths:
                cell.width = Inches(col_widths[i])
    doc.add_paragraph()
    return table


def add_heading(doc, text, size=16, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color


def add_subheading(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(0x2E, 0x5C, 0x8A)


def add_note(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GRAY


def build_reliability_section(doc, reliability):
    add_heading(doc, "1. Reliability Statistics", size=14)
    add_note(doc, "Mean, median, standard deviation, and coefficient of variation across the 5 scored "
                  "runs of each stochastic-sample prompt. Lower coefficient of variation = more "
                  "consistent scoring across repeated calls.")

    for provider, pdata in reliability.get("providers", {}).items():
        add_subheading(doc, provider.capitalize())
        avg_cv = pdata.get("average_coefficient_of_variation")
        add_note(doc, f"Average CV across {pdata.get('n_prompts_analyzed', 0)} analyzed prompts: {avg_cv}")

        rows = []
        for pid, stats in sorted(pdata.get("per_prompt", {}).items()):
            cv = stats.get("coefficient_of_variation")
            highlight = None
            if cv is not None:
                highlight = "green" if cv < 0.15 else ("red" if cv > 0.35 else None)
            rows.append([
                (pid, None), (stats.get("mean"), None), (stats.get("median"), None),
                (stats.get("standard_deviation"), None), (cv, highlight),
            ])
        if rows:
            _add_table(doc, ["Prompt ID", "Mean", "Median", "Std Dev", "Coeff. of Variation"], rows)


def build_ci_section(doc, analysis):
    add_heading(doc, "2. Confidence Intervals (95%)", size=14)
    add_note(doc, "95% confidence interval on the mean total score, per provider.")

    rows = []
    for provider, ci in analysis.get("confidence_intervals", {}).items():
        if ci:
            rows.append([
                (provider.capitalize(), None), (ci["mean"], None),
                (f"[{ci['ci_lower']}, {ci['ci_upper']}]", None), (ci["n"], None),
            ])
    _add_table(doc, ["Provider", "Mean Score", "95% CI", "N"], rows)


def build_hypothesis_section(doc, analysis):
    add_heading(doc, "3. Hypothesis Testing", size=14)

    tests = analysis.get("hypothesis_tests", {})

    add_subheading(doc, "H1: Multi-step Reasoning vs. Knowledge Retrieval")
    rows = []
    for provider, result in tests.get("H1_reasoning_vs_retrieval", {}).items():
        t = result.get("t_test")
        if t:
            sig = "Yes" if t["significant_at_0.05"] else "No"
            highlight = "green" if t["significant_at_0.05"] else None
            rows.append([
                (provider.capitalize(), None), (t["mean_a"], None), (t["mean_b"], None),
                (t["p_value"], None), (sig, highlight), (result.get("cohens_d"), None),
            ])
    _add_table(doc, ["Provider", "Reasoning Mean", "Retrieval Mean", "p-value", "Significant?", "Cohen's d"], rows)

    add_subheading(doc, "H2: Instruction Following Score vs. Difficulty (Correlation)")
    rows = []
    for provider, result in tests.get("H2_instruction_following_vs_difficulty", {}).items():
        rows.append([(provider.capitalize(), None), (result.get("correlation_difficulty_vs_score"), None), (result.get("n"), None)])
    _add_table(doc, ["Provider", "Correlation (Difficulty vs Score)", "N"], rows)

    add_subheading(doc, "H3: Hallucination Stress Test - Provider Comparisons")
    rows = []
    for comparison, result in tests.get("H3_hallucination_provider_comparison", {}).items():
        t = result.get("t_test")
        if t:
            sig = "Yes" if t["significant_at_0.05"] else "No"
            highlight = "green" if t["significant_at_0.05"] else None
            rows.append([
                (comparison.replace("_", " "), None), (t["p_value"], None),
                (sig, highlight), (result.get("cohens_d"), None),
            ])
    _add_table(doc, ["Comparison", "p-value", "Significant?", "Cohen's d"], rows)


def build_correlation_regression_section(doc, analysis):
    add_heading(doc, "4. Correlation and Regression", size=14)

    add_subheading(doc, "Correlation: Difficulty and Latency vs. Score")
    rows = []
    for provider, corr in analysis.get("correlation", {}).items():
        rows.append([
            (provider.capitalize(), None), (corr.get("difficulty_vs_score"), None), (corr.get("latency_vs_score"), None),
        ])
    _add_table(doc, ["Provider", "Corr(Difficulty, Score)", "Corr(Latency, Score)"], rows)

    add_subheading(doc, "Regression: Predicting Score from Difficulty")
    rows = []
    for provider, reg in analysis.get("regression", {}).items():
        if reg:
            rows.append([
                (provider.capitalize(), None), (reg.get("intercept"), None),
                (reg.get("slope"), None), (reg.get("r_squared"), None),
            ])
    _add_table(doc, ["Provider", "Intercept", "Slope", "R-squared"], rows)


def build_inter_rater_section(doc, analysis):
    inter_rater_data = analysis.get("inter_rater_agreement")
    if not inter_rater_data:
        return

    add_heading(doc, "5. Inter-Rater Agreement", size=14)
    add_note(doc, "From Week 6's human validation. Cohen's Kappa corrects for agreement expected by "
                  "chance alone; the weighted variant additionally accounts for the ordinal 0-5 scale "
                  "(a 0-vs-5 miss is worse than a 3-vs-4 miss).")

    add_subheading(doc, "Human Raters vs. LLM Judge")
    rows = []
    for rater, result in inter_rater_data.get("human_vs_judge", {}).items():
        kappa = result.get("cohens_kappa")
        highlight = "green" if kappa is not None and kappa >= 0.4 else ("red" if kappa is not None and kappa < 0.2 else None)
        rows.append([
            (rater.replace("rater", "Rater "), None), (result.get("n_criteria_compared"), None),
            (f"{result.get('exact_match_rate', 0)*100:.1f}%", None),
            (f"{result.get('within_1_point_rate', 0)*100:.1f}%", None),
            (result.get("pearson_correlation"), None), (kappa, highlight),
            (result.get("cohens_kappa_quadratic_weighted"), None),
        ])
    _add_table(doc, ["Rater", "N", "Exact Match", "Within 1pt", "Correlation", "Kappa", "Weighted Kappa"], rows)

    add_subheading(doc, "Inter-Rater Agreement (Human vs. Human)")
    rows = []
    for pair, result in inter_rater_data.get("inter_rater", {}).items():
        kappa = result.get("cohens_kappa")
        highlight = "green" if kappa is not None and kappa >= 0.4 else ("red" if kappa is not None and kappa < 0.2 else None)
        rows.append([
            (pair.replace("_", " ").title(), None), (result.get("n_criteria_compared"), None),
            (f"{result.get('exact_match_rate', 0)*100:.1f}%", None),
            (result.get("pearson_correlation"), None), (kappa, highlight),
            (result.get("cohens_kappa_quadratic_weighted"), None),
        ])
    _add_table(doc, ["Comparison", "N", "Exact Match", "Correlation", "Kappa", "Weighted Kappa"], rows)


def main():
    reliability = _load(RELIABILITY_PATH)
    analysis = _load(ANALYSIS_PATH)

    if reliability is None and analysis is None:
        print("FAILURE: neither reliability_statistics.json nor statistical_analysis.json found. "
              "Run compute_reliability_statistics.py and statistical_analysis.py first.")
        return

    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Week 7: Statistical Analysis Report")
    r.bold = True
    r.font.size = Pt(20)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = subtitle.add_run("Reliability, Confidence Intervals, Hypothesis Testing, Correlation, and Regression")
    r2.italic = True
    r2.font.size = Pt(11)
    r2.font.color.rgb = GRAY
    doc.add_paragraph()

    if reliability:
        build_reliability_section(doc, reliability)
    if analysis:
        build_ci_section(doc, analysis)
        build_hypothesis_section(doc, analysis)
        build_correlation_regression_section(doc, analysis)
        build_inter_rater_section(doc, analysis)

    os.makedirs(STATS_DIR, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Report written to: {os.path.abspath(OUTPUT_PATH)}")


if __name__ == "__main__":
    main()

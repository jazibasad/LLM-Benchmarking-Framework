"""
build_results_docx.py

Creates one human-readable Word document per provider, showing every prompt
alongside its model response(s) - satisfying the requirement that the raw
data (prompts and their corresponding answers) be submitted as separate,
structured, human-readable files, not just raw JSON.

Reads from 05_Logs_Results/Combined_Results/<provider>_prompts_and_results.json
(built by build_results_tables.py), so run that script first.

STOCHASTIC-SAMPLE PROMPTS: for the 20 prompts repeated 5 times each, ALL 5
response variants are shown together in the same row, clearly labeled
"Run 1:", "Run 2:", etc., so a reader can see the actual run-to-run
variation directly - this is the point of the reliability sample, and
hiding it behind only the first run would defeat that purpose. Standard
single-run prompts show just one response, as before.

OUTPUT FILES (written to 05_Logs_Results/Readable_Results/):
  - gemini_results.docx
  - mistral_results.docx
  - groq_results.docx

Requires: pip install python-docx

Run with the VS Code Run button:
    python 03_Code/build_results_docx.py
"""

import json
import os

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = os.path.join(os.path.dirname(__file__), "..")
COMBINED_RESULTS_DIR = os.path.join(BASE_DIR, "05_Logs_Results", "Combined_Results")
OUTPUT_DIR = os.path.join(BASE_DIR, "05_Logs_Results", "Readable_Results")

PROVIDERS = {
    "gemini": {
        "input_file": os.path.join(COMBINED_RESULTS_DIR, "gemini_prompts_and_results.json"),
        "output_name": "gemini_results.docx",
        "title": "Gemini - Prompts and Results",
    },
    "mistral": {
        "input_file": os.path.join(COMBINED_RESULTS_DIR, "mistral_prompts_and_results.json"),
        "output_name": "mistral_results.docx",
        "title": "Mistral - Prompts and Results",
    },
    "groq": {
        "input_file": os.path.join(COMBINED_RESULTS_DIR, "groq_prompts_and_results.json"),
        "output_name": "groq_results.docx",
        "title": "Groq - Prompts and Results",
    },
}

COLUMNS = [
    ("ID", 0.5),
    ("Category", 1.0),
    ("Difficulty", 0.6),
    ("Prompt", 2.0),
    ("Response(s)", 3.0),
    ("Temperature", 0.6),
    ("Latency (s)", 0.9),
    ("Runs", 0.4),
]


def _load_records(input_file):
    if not os.path.exists(input_file):
        return None
    with open(input_file, "r", encoding="utf-8") as f:
        return json.load(f)


def _format_responses(entry) -> str:
    """Response text only - for single-run prompts, just the text. For
    multi-run prompts, each run's text labeled by run number, aligned with
    the separate Temperature/Latency columns' own per-run breakdown."""
    runs = entry.get("runs", [])
    if not runs:
        return "(no response recorded)"

    if len(runs) == 1:
        return runs[0].get("response_text", "(no response recorded)")

    parts = []
    for run in runs:
        run_num = run.get("run_number", "?")
        text = run.get("response_text", "(no response recorded)")
        parts.append(f"--- Run {run_num} ---\n{text}")
    return "\n\n".join(parts)


def _format_temperature_column(entry) -> str:
    """Temperature is fixed per call, but shown per-run row-for-row so it
    lines up with the Response(s) and Latency columns, since the table is
    the single source of truth for what was actually recorded for each run."""
    runs = entry.get("runs", [])
    if not runs:
        return "-"
    if len(runs) == 1:
        temp = runs[0].get("temperature")
        return str(temp) if temp is not None else "-"
    return "\n\n".join(str(r.get("temperature", "-")) for r in runs)


def _format_latency_column(entry) -> str:
    """Latency measured per API call - genuinely varies run to run, so each
    run's latency is shown on its own line, aligned with the corresponding
    run in the Response(s) column."""
    runs = entry.get("runs", [])
    if not runs:
        return "-"
    if len(runs) == 1:
        latency = runs[0].get("latency_seconds")
        return str(latency) if latency is not None else "-"
    return "\n\n".join(str(r.get("latency_seconds", "-")) for r in runs)


def _set_cell_shading(cell, hex_color):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _set_repeat_header(row):
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def _add_multiline_text(paragraph, text, font_size, bold=False, bold_lines_starting_with=None):
    """Real, visible line breaks for multi-line content. Optionally makes
    lines starting with a given prefix (e.g. '--- Run') bold, so the run
    labels stand out visually from the response text itself."""
    lines = str(text).split("\n")
    for i, line in enumerate(lines):
        if i > 0:
            paragraph.add_run().add_break(WD_BREAK.LINE)
        is_label = bold_lines_starting_with and line.startswith(bold_lines_starting_with)
        run = paragraph.add_run(line)
        run.font.size = Pt(font_size)
        run.bold = bold or is_label
        if is_label:
            run.font.color.rgb = RGBColor(0x2E, 0x5C, 0x8A)


def build_doc_for_provider(name, config):
    records = _load_records(config["input_file"])

    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width, section.page_height = new_width, new_height
    section.left_margin = Inches(0.4)
    section.right_margin = Inches(0.4)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(config["title"])
    run.bold = True
    run.font.size = Pt(18)

    if records is None:
        warn = doc.add_paragraph()
        warn.alignment = WD_ALIGN_PARAGRAPH.CENTER
        warn_run = warn.add_run(
            f"No consolidated results found at {config['input_file']}. "
            f"Run build_results_tables.py first."
        )
        warn_run.italic = True
        warn_run.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)
        os.makedirs(OUTPUT_DIR, exist_ok=True)
        doc.save(os.path.join(OUTPUT_DIR, config["output_name"]))
        print(f"[{name}] No input data - placeholder document written.")
        return 0

    sampled_count = sum(1 for r in records if r.get("stochastic_sample"))
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run(
        f"{len(records)} prompts - {sampled_count} stochastic-sample prompts show all 5 response runs, "
        f"each with its own temperature and latency"
    )
    sub_run.italic = True
    sub_run.font.size = Pt(10)
    sub_run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    table = doc.add_table(rows=1, cols=len(COLUMNS))
    table.style = "Table Grid"
    table.autofit = False

    header_row = table.rows[0]
    _set_repeat_header(header_row)
    for i, (col_name, width) in enumerate(COLUMNS):
        cell = header_row.cells[i]
        cell.width = Inches(width)
        p = cell.paragraphs[0]
        r = p.add_run(col_name)
        r.bold = True
        r.font.size = Pt(9)
        _set_cell_shading(cell, "D9E2F3")

    for entry in records:
        row = table.add_row()
        is_sampled = entry.get("stochastic_sample", False)
        num_runs = len(entry.get("runs", []))

        values = [
            entry.get("id", ""),
            entry.get("category", ""),
            entry.get("difficulty", ""),
            entry.get("prompt", ""),
            _format_responses(entry),
            _format_temperature_column(entry),
            _format_latency_column(entry),
            str(num_runs),
        ]
        for i, (col_name, width) in enumerate(COLUMNS):
            cell = row.cells[i]
            cell.width = Inches(width)
            p = cell.paragraphs[0]
            if col_name == "Response(s)":
                _add_multiline_text(p, values[i], font_size=8, bold_lines_starting_with="--- Run")
            else:
                _add_multiline_text(p, values[i], font_size=8.5)
            if is_sampled:
                _set_cell_shading(cell, "FFF3CD")

    os.makedirs(OUTPUT_DIR, exist_ok=True)
    output_path = os.path.join(OUTPUT_DIR, config["output_name"])
    doc.save(output_path)

    print(f"[{name}] {len(records)} rows written to {output_path} ({sampled_count} highlighted multi-run prompts)")
    return len(records)


def main():
    print(f"Building human-readable prompt+result tables in: {os.path.abspath(OUTPUT_DIR)}\n")
    for name, config in PROVIDERS.items():
        build_doc_for_provider(name, config)
    print("\nDone.")


if __name__ == "__main__":
    main()
